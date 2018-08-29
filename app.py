# -*- coding: utf-8 -*-
"""
Created on Tue Aug 28 23:31:14 2018

@author: udays
"""



from flask import Flask, render_template,request,redirect, url_for
from docx.api import Document
import sqlite3 as lite
from flask_wtf.csrf import CSRFProtect, CSRFError
import re
from newinput import input_form, Delete_form, updqate_form

app = Flask(__name__)
app.config.update(dict(
    SECRET_KEY="powerful secretkey",
    WTF_CSRF_SECRET_KEY="a csrf secret key"
))


# function to inport data into database!! only once used
def import_data(filename):
    conn = lite.connect("input.db")
    cur = conn.cursor()
    cur.execute("create table IF NOT EXISTS input('id' integer, 'description' text, 'datetime' integer, 'longitute' real, 'latitude' real, 'elevation' integer)")
    cur.execute("select count(*) from input")
    count=cur.fetchall()
    if(count[0][0]==0):
        i=1
        document = Document(filename)
        paragraph = document.paragraphs[0]
        while(len(paragraph.text)!=0):
            document = Document(filename)
            paragraph = document.paragraphs[i]
            col = re.split(r'\t+', paragraph.text)
            cur.execute("insert into input values(?,?,datetime(?),?,?,?)",(col[0],col[1],col[2],col[3],col[4],col[5]))
            conn.commit()
            i=i+1
    conn.close()
    
    
# function to fetch all data from database
def display():
    conn = lite.connect("input.db")
    cur = conn.cursor()
    cur.execute("select * from input")
    inputdata = cur.fetchall()
    conn.close()
    print(inputdata)
    return inputdata
 
# function to insert a new record
def insert(col1):
    conn = lite.connect("input.db")
    cur = conn.cursor()
    cur.execute("insert into input values(?,?,datetime(?),?,?,?)",(col1[0],col1[1],col1[2],col1[3],col1[4],col1[5]))
    conn.commit()
    conn.close()

# function to delet a record according to the index    
def delete(index1):
    inputdata =[]
    conn = lite.connect("input.db")
    cur = conn.cursor()
    cur.execute("select * from input")
    inputdata = cur.fetchall()
    for i in range (0,len(inputdata)):
        
        if(i==index1-1):
            temparray =[]
            temparray.append(inputdata[i][0])
            temparray.append(inputdata[i][1])
            temparray.append(inputdata[i][2])
            temparray.append(inputdata[i][3])
            temparray.append(inputdata[i][4])
            temparray.append(inputdata[i][5])
    if len(temparray)==0:
        conn.close()
        return False
        
    else:
        cur.execute("delete from input where id=? and description =? and datetime = ? and longitute=? and latitude =? and elevation =?",(temparray[0],temparray[1],temparray[2],temparray[3],temparray[4],temparray[5]))
        conn.commit()
        conn.close()
        return True
        

#function to update a record in database 
def updateData(col2):
    inputdata =[]
    conn = lite.connect("input.db")
    cur = conn.cursor()
    cur.execute("select * from input")
    inputdata = cur.fetchall()
    for i in range (0,len(inputdata)):
        
        if(i==col2[0]-1):
            temparray =[]
            temparray.append(inputdata[i][0])
            temparray.append(inputdata[i][1])
            temparray.append(inputdata[i][2])
            temparray.append(inputdata[i][3])
            temparray.append(inputdata[i][4])
            temparray.append(inputdata[i][5])
    if len(temparray)==0:
        conn.close()
        return False
        
    else:
        cur.execute("update input set id =?,description =?,datetime = ?,longitute=?,latitude =?,elevation =? where id=? and description =? and datetime = ? and longitute=? and latitude =? and elevation =?",(col2[1],col2[2],col2[3],col2[4],col2[5],col2[6],temparray[0],temparray[1],temparray[2],temparray[3],temparray[4],temparray[5]))
        conn.commit()
        conn.close()
        return True

# app routes to perform 4 functions display. add, delete and update
@app.route('/')
def main():
    return redirect(url_for('index'))

@app.route('/index')
def index():
    inputdata = display()
    return render_template('index.html', data = inputdata)

@app.route('/add', methods = ['GET','POST'])
def add():
    form1 = input_form()
    if request.method == 'POST' and form1.validate():
        new_array = []
        new_array.append(form1.id.data)
        new_array.append(form1.description.data)
        new_array.append(form1.datetime.data)
        new_array.append(form1.Longitude.data)
        new_array.append(form1.Latitude.data)
        new_array.append(form1.Elevation.data)
        if form1.validate_on_submit():
            insert(new_array)
            return redirect(url_for('index'))
    return render_template('add.html', form = form1)

@app.route('/deleteR', methods = ['GET','POST'])
def deleteR():
    form = Delete_form()
    if request.method == 'POST' and form.validate_on_submit():
        result = delete(form.Index.data)
        if(result==False):
            return '<h1> Wrong Index </h1>'
        else:
            return redirect(url_for('index'))
    return render_template('deleteR.html', form = form)

@app.route('/update', methods = ['GET','POST'])
def update():
    form2 = updqate_form()
    new_array = []
    new_array.append(form2.Index.data)
    new_array.append(form2.id.data)
    new_array.append(form2.description.data)
    new_array.append(form2.datetime.data)
    new_array.append(form2.Longitude.data)
    new_array.append(form2.Latitude.data)
    new_array.append(form2.Elevation.data)
    if request.method == 'POST' and form2.validate_on_submit():
        result = updateData(new_array)
        if result==False:
            return '<h1> Wrong Index </h1>'
        else:
            return redirect(url_for('index'))
    return render_template('update.html', form = form2)
        

if __name__ == '__main__':
#    import_data("input.txt.docx")
    app.run(debug=True)